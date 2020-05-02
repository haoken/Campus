# !/usr/bin/env python
# coding: utf-8

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import  Pt
import math

def printDelegationTable(document,delegation=[[],[]],count=1,dict = {}):
    ######男生分组
    if len(delegation[0])!=0:
        count_offset = count - 1
        number_delegation = len(delegation[0])
        cols = 4#每行又4个单元格
        rows = math.ceil(len(delegation[0])/float(cols))
        document.add_paragraph('男子组')
        table = document.add_table(1, cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells = table.rows[0].cells
        for first_row_cell in row_cells:
            if count-count_offset<=number_delegation:
                dict[delegation[0][count-count_offset-1]]= count
                first_row_cell.text=str(count)+'  '+delegation[0][count-count_offset-1]
                count = count + 1
        for i in range(1, rows, 1):
            row_cells = table.add_row().cells
            for j in range(0, cols, 1):
                if count-count_offset <= number_delegation:
                    dict[delegation[0][count - count_offset - 1]] = count
                    row_cells[j].text = str(count)+'  '+delegation[0][count-count_offset-1]
                    count = count + 1
    #####女生分组
    if len(delegation[1])!=0:
        count_offset = count-1
        cols = 4
        number_delegation = len(delegation[1])
        rows = math.ceil(number_delegation/float(cols))
        document.add_paragraph('女子组')
        table = document.add_table(1, cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells = table.rows[0].cells
        for first_row_cell in row_cells:
            if count-count_offset <= number_delegation:
                dict[delegation[1][count - count_offset - 1]] = count
                first_row_cell.text = str(count) + '  ' + delegation[1][count- count_offset - 1]
                count = count + 1
        for i in range(1, rows, 1):
            row_cells = table.add_row().cells
            for j in range(0, cols, 1):
                if count-count_offset <= number_delegation:
                    dict[delegation[1][count - count_offset - 1]] = count
                    row_cells[j].text = str(count) + '  ' + delegation[1][count- count_offset - 1]
                    count = count + 1
    return count



#hostunit:主办单位str
#orginizer:承办单位str
#time:举办时间str
#place:举办位置str
#itemlist[]:项目表
#delegation:代表团名单（9个代表团的参赛名单）
#  delegation[n]:是第n个代表团的名单
#  delegation[n][0]:第n个代表团名单的男子名单
#  delegation[n][1]:第n个代表团名单的女子名单
#  delegation[n][0][0]:第n个代表团名单的男子名单中的第一人
#   ...
#itemgrouplist:项目分组表
#   itemgrouplists[0]:第一个项目的分组情况
#   itemgrouplists[0][0]:第一个项目的男子分组
#   itemgrouplists[0][0][0]:第一个项目的男子分组的第一个分组
#   itemgrouplists[0][0][0][0]:第一个项目的男子分组的第一个分组的第一个运动员
#   ...
#itemName:与上面项目分组表相对应的项目名称表
#   itemName[0]:第一个项目的名称
#   itemName[1]:第二个项目的名称
#   ...
#
def writetodocx(hostunit,organizer,time,place,itemlist=[],
                delegation=[[[],[]],[[],[]],[[],[]],[[],[]],[[],[]],[[],[]],[[],[]],[[],[]],[[],[]]],
                itemgrouplist=[[[[]]]],itemName = [],):
    print(hostunit)
    print(organizer)
    print(time)
    print(place)
    document = Document()
    head = document.add_heading('校运会秩序册',level=0)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #######################################################
    head = document.add_heading('1.竞赛规程',level=1,)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_heading('一、主办单位',level=2)
    document.add_paragraph(hostunit)

    document.add_heading('二、承办单位',level=2)
    document.add_paragraph(organizer)

    document.add_heading('三、时间地点',level=2)
    document.add_paragraph(time)
    document.add_paragraph(place)

    document.add_heading('四、竞赛项目',level=2)
    paragraph = document.add_paragraph(itemlist[0]+' , ')
    for item in itemlist[1:]:
        paragraph.add_run(item+' , ')
    document.add_page_break()
    ###################################################
    dict_athlete_order = {}#该字典 存放所有的运动员和他们对应的编号
    head = document.add_heading('2.代表队名单',level=1,)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    head = document.add_heading('(1)知行书院', level=3, )
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count = printDelegationTable(document,delegation[0],1,dict_athlete_order)
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    head = document.add_heading('(2)思源书院', level=3 )
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count = printDelegationTable(document,delegation[1],count,dict_athlete_order)
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    head = document.add_heading('(3)弘毅书院', level=3, )
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count = printDelegationTable(document,delegation[2],count,dict_athlete_order)
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    head = document.add_heading('(4)敬一书院', level=3, )
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count = printDelegationTable(document,delegation[3],count,dict_athlete_order)
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    head = document.add_heading('(5)至诚书院', level=3, )
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count = printDelegationTable(document,delegation[4],count,dict_athlete_order)
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    head = document.add_heading('(6)修远书院', level=3, )
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count = printDelegationTable(document,delegation[5],count,dict_athlete_order)
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    head = document.add_heading('(7)明德书院', level=3, )
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count = printDelegationTable(document,delegation[6],count,dict_athlete_order)
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    head = document.add_heading('(8)德馨书院', level=3, )
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count = printDelegationTable(document,delegation[7],count,dict_athlete_order)
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    head = document.add_heading('(9)研究生院', level=3, )
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count = printDelegationTable(document,delegation[8],count,dict_athlete_order)
    document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
    document.add_page_break()
###########################################################
    head = document.add_heading('3.竞赛日程',level=1,)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_page_break()

###########################################################
    dict_number = {1:'一',2:'二',3:'三',4:'四',5:'五',6:'六',7:'七',8:'八'}
    head = document.add_heading('4.项目分组表',level=1,)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count = 0
    for itemgroup in itemgrouplist: ##itemgroup 是单个分组情况
        document.add_paragraph()
        if len(itemgroup[0][0])!=0:#男子组的第一组有人
            paragraph = document.add_paragraph()
            paragraph.add_run('男子组 '+itemname[count]).bold = True
            malegroup = itemgroup[0]
            groupnumber = 1
            for group in malegroup:
                document.add_paragraph('第'+str(groupnumber)+'组:')
                groupnumber = groupnumber + 1
                table = document.add_table(rows=1,cols=8)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table_cells = table.rows[0].cells
                count_cells = 1
                for cell in table_cells[0:len(group)]:
                    cell.text = '     '+dict_number[count_cells]+'\n     '+str(dict_athlete_order[group[count_cells-1]])+'\n'+group[count_cells-1]
                    count_cells = count_cells+1
        document.add_paragraph()#隔一行
        if len(itemgroup[1][0])!=0:#女子组的第一组有人
            paragraph = document.add_paragraph()
            paragraph.add_run('女子组 ' + itemname[count]).bold = True
            femalegroup = itemgroup[1]
            groupnumber = 1
            for group in femalegroup:
                document.add_paragraph('第' + str(groupnumber) + '组:')
                groupnumber = groupnumber + 1
                table = document.add_table(rows=1, cols=8)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table_cells = table.rows[0].cells
                count_cells = 1
                for cell in table_cells[0:len(group)]:
                    cell.text = '     ' + dict_number[count_cells] + '\n     ' + str(
                        dict_athlete_order[group[count_cells - 1]]) + '\n' + group[count_cells - 1]
                    count_cells = count_cells + 1
        count = count+1
    document.add_page_break()
##########################################################
    document.save('校运会秩序册.docx')
    return 1


# delegation = [[['李嘉浩1','李嘉浩2'],['李嘉浩3','李嘉浩4','李嘉浩5']],
#               [['博尔特1','博尔特2','博尔特3'],['博尔特4','博尔特5']],
#               [['刘翔1'],['刘翔2','刘翔3','刘翔4','刘翔5','刘翔6']],
#               [['岳云鹏1','岳云鹏2','岳云鹏3'],['岳云鹏4','岳云鹏5']],
#               [['林峰1','林峰2','林峰3','林峰4','林峰5'],['林峰6']],
#               [['小龟1','小龟2','小龟3'],['小龟4','小龟5','小龟6']],
#               [['高键铧1','高键铧2','高键铧3'],['林峰6','林峰7','林峰8']],[['林峰9','林峰10','林峰11'],['张继峰1','张继峰2','张继峰3']],[['张智成1','张智成2','张智成3'],['张智成4','张智成5','张智成6','张智成7']]]
# itemgroup = [
#     [[['高键铧1','高键铧2','高键铧3'],['小龟1','小龟2'],['张智成1','张智成2']],[['林峰5','林峰6'],['岳云鹏4','岳云鹏5'],['刘翔4','刘翔5','刘翔6']]],
#     [[['博尔特1','博尔特2','博尔特3'],['李嘉浩3','李嘉浩4','李嘉浩5']],[['林峰3','林峰4','林峰5'],['张继峰1','张继峰2','张继峰3']]],
#     [[['张智成1','张智成2','张智成3'],['张智成4','张智成5','张智成6','张智成7']],[['林峰6','林峰7','林峰8'],['林峰9','林峰10','林峰11']]]
# ]
# itemname = ['跑步（100米）','跑步（400米）','跑步（800米）']
#
# itemlist = ['跑步（100米）','跑步（200米）','跑步（400米）']
# writetodocx('汕头大学体育委员会','汕头大学体育部','2020.9.12~2020.9.13','汕头大学田径场',itemlist,delegation,itemgroup,itemname)

